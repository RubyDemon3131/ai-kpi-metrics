from openai import OpenAI
import requests
import pdfplumber
from io import BytesIO
import re
from docx import Document

client = OpenAI(api_key='Your API KEY')

def extract_text_from_url(url):
    print("Extracting text from the Given URL......")
    response = requests.get(url)
    content_type = response.headers.get('content-type')
    if 'pdf' in content_type:
        return extract_text_from_pdf(response.content)
    else:
        return response.text

def extract_text_from_pdf(pdf_content):
    print("Extracting text from the Given PDF......")
    text = ""
    with pdfplumber.open(BytesIO(pdf_content)) as pdf:
        for page in pdf.pages:
            text += page.extract_text()
    return text

def extract_kpis(report_text):
    print("Extracting KPIs from the given text......")
    kpi_patterns = {
        'sales': r'Sales: \$(\d+[\d,]*)',
        'conversion_rate': r'Conversion Rate: (\d+\.?\d*)%',
        'customer_acquisition_cost': r'Customer Acquisition Cost: \$(\d+[\d,]*)'
    }
    kpis = {}
    for kpi, pattern in kpi_patterns.items():
        match = re.search(pattern, report_text)
        if match:
            kpis[kpi] = match.group(1)
    return kpis

def generate_summary(report_text, kpi_summary, notes_text, tone):
    print(f"Sending the request to GPT4 to write summary in {tone} tone......")
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "user", "content": f"Here is a marketing report and some notes. Summarize the key points, including KPI metrics, and provide the summary in a {tone} tone. Report: {report_text} Notes: {notes_text} KPIs: {kpi_summary}."}
        ],
        max_tokens=1500,
        temperature=0.5,
    )
    print("Completed!")
    return response.choices[0].message.content

def create_word_document(content, tone):
    doc = Document()
    doc.add_heading(f'Summary ({tone.capitalize()} Tone)', 0)

    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)

    file_name = f"summary_{tone}.docx"
    doc.save(file_name)
    return file_name

def main():
    report_url = 'https://www.newessays.co.uk/wp-content/uploads/2011/12/Businessreport-Sample.pdf'
    notes = """
    - **Executive Summary**: Overview of company's market position and financial health, including recent strategic decisions.
    - **Market Analysis**: Trends and consumer behavior, competitor market shares, and potential opportunities and threats.
    - **Sales Performance**: Sales figures by product line and region, year-over-year growth, and sales strategy effectiveness.
    - **Marketing Strategies**: Recent campaigns, their outcomes, and ROI analysis of marketing initiatives.
    - **Customer Insights**: Demographics, satisfaction survey results, retention rates, and loyalty programs.
    - **Financial Overview**: Summary of revenue, expenses, profit margins, and financial projections.
    - **Key Performance Indicators (KPIs)**: Sales conversion rates, customer acquisition cost (CAC), lifetime value (LTV), and campaign performance metrics.
    - **Strategic Recommendations**: Suggestions for market penetration, marketing strategy optimization, and potential investment areas.
    """

    tones = ['friendly', 'professional', 'enthusiastic']

    report_text = extract_text_from_url(report_url)
    print("Completed!")
    
    kpis = extract_kpis(report_text)
    kpi_summary = ", ".join([f"{k}: {v}" for k, v in kpis.items()])
    print("Completed!")
    
    summaries = {tone: generate_summary(report_text, kpi_summary, notes, tone) for tone in tones}

    print("Saving the documents for each tone.....")
    for tone, summary in summaries.items():
        create_word_document(summary, tone)
    print("Completed!")

if __name__ == "__main__":
    main()